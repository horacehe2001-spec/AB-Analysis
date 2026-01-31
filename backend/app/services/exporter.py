from __future__ import annotations

import base64
import io
import logging
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from app.db.models import SessionModel

logger = logging.getLogger(__name__)

_EFFECT_LEVEL_CN = {"small": "小", "medium": "中等", "large": "大"}
_EFFECT_TYPE_CN = {
    "cohens_d": "Cohen's d",
    "r_squared": "R²",
    "eta_squared": "η²",
    "cramers_v": "Cramér's V",
}


def _to_iso(dt: datetime) -> str:
    if not dt.tzinfo:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")


def _collect_analyses(session: SessionModel) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for m in session.messages:
        if m.role == "assistant" and m.analysis:
            results.append(m.analysis)
    return results


# ---------------------------------------------------------------------------
# LLM conclusion
# ---------------------------------------------------------------------------

def _generate_llm_conclusion(
    analyses: list[dict[str, Any]],
    data_summary: dict[str, Any] | None,
    llm_config: dict | None,
) -> str | None:
    if not llm_config or not llm_config.get("api_key"):
        return None
    try:
        from app.services.llm.client import LLMModelConfig, call_llm_json

        summaries = []
        for i, a in enumerate(analyses, 1):
            es = a.get("effect_size", {})
            summaries.append(
                f"分析{i}: 方法={a.get('method_name','未知')}, "
                f"p值={a.get('p_value')}, "
                f"显著={'是' if a.get('significant') else '否'}, "
                f"效应量={_EFFECT_TYPE_CN.get(es.get('type',''), es.get('type',''))}="
                f"{es.get('value')}({_EFFECT_LEVEL_CN.get(es.get('level',''), es.get('level',''))}), "
                f"解读={a.get('interpretation','')}"
            )

        data_info = ""
        if data_summary:
            data_info = f"数据：{data_summary.get('rows')}行×{data_summary.get('columns')}列"

        system_prompt = (
            "你是资深统计分析师。根据假设检验结果撰写分析结论。\n"
            "要求：1.总结核心发现 2.解释业务含义 3.给出可操作建议 "
            "4.多组分析时做综合对比 5.语言专业但易懂 6.直接输出文本，不要JSON"
        )
        user_prompt = f"{data_info}\n\n分析结果：\n" + "\n".join(summaries) + "\n\n请撰写分析结论。"

        cfg_dict = {**llm_config, "max_tokens": max(int(llm_config.get("max_tokens", 2048)), 2048)}
        cfg = LLMModelConfig(**cfg_dict)
        return call_llm_json(config=cfg, system_prompt=system_prompt, user_prompt=user_prompt).strip()
    except Exception as e:
        logger.warning("LLM conclusion generation failed: %s", e)
        return None


# ---------------------------------------------------------------------------
# Shared: build per-analysis process text
# ---------------------------------------------------------------------------

def _analysis_process_md(a: dict[str, Any], index: int, data_summary: dict[str, Any] | None) -> list[str]:
    """Return markdown lines for one analysis's 6-step process."""
    lines: list[str] = []
    method_name = a.get("method_name", a.get("method", "未知方法"))
    p_value = a.get("p_value")
    significant = a.get("significant", False)
    es = a.get("effect_size", {})
    interpretation = a.get("interpretation", "")
    suggestions = a.get("suggestions", [])

    lines.append(f"### 分析 {index}：{method_name}")
    lines.append("")

    # 1 数据识别
    lines.append("**步骤一 · 数据识别**")
    if data_summary:
        cols = data_summary.get("column_names", [])
        col_preview = "、".join(cols[:12])
        if len(cols) > 12:
            col_preview += " ..."
        lines.append(f"- 样本量：{data_summary.get('rows', '-') } 行")
        lines.append(f"- 变量数：{data_summary.get('columns', '-') } 列")
        if cols:
            lines.append(f"- 变量列表：{col_preview}")
    lines.append(f"- 分析类型：{method_name}")
    lines.append("")

    # 2 前提条件校验
    lines.append("**步骤二 · 前提条件校验**")
    param_methods = {"t_test", "anova", "pearson", "linear_regression"}
    if a.get("method") in param_methods:
        lines.append("- 前提假设：数据近似正态分布")
        lines.append("- 检验方式：参数检验")
    else:
        lines.append("- 前提假设：无分布假设要求")
        lines.append("- 检验方式：非参数/分布无关检验")
    lines.append("- 显著性水平 α = 0.05")
    lines.append("")

    # 3 方法选择
    lines.append("**步骤三 · 方法选择**")
    lines.append(f"- 采用方法：{method_name}")
    decision_hint = {
        "t_test": "两组独立样本 → 连续变量 → 独立样本 t 检验",
        "mann_whitney_u": "两组独立样本 → 非正态 → Mann–Whitney U",
        "anova": "多组比较 → 连续变量 → 单因素 ANOVA",
        "kruskal": "多组比较 → 非正态 → Kruskal–Wallis",
        "chi_square": "分类变量 → 频数数据 → 卡方检验",
        "pearson": "连续变量 → 线性相关 → Pearson 相关",
        "spearman": "变量有序/非正态 → Spearman 相关",
        "linear_regression": "预测关系 → 线性回归",
    }.get(a.get("method") or "", None)
    if decision_hint:
        lines.append(f"- 决策路径：{decision_hint}")
    lines.append("")

    # 4 统计计算
    lines.append("**步骤四 · 统计检验**")
    if p_value is not None:
        lines.append(f"- p = {p_value:.6g}")
        if significant:
            lines.append("- **p < 0.05 → 拒绝零假设，差异 / 关系具有统计显著性**")
        else:
            lines.append("- p ≥ 0.05 → 无法拒绝零假设，差异 / 关系不显著")
    lines.append("")

    # 5 效应量
    lines.append("**步骤五 · 效应量分析**")
    es_type = _EFFECT_TYPE_CN.get(es.get("type", ""), es.get("type", ""))
    es_value = es.get("value")
    es_level = _EFFECT_LEVEL_CN.get(es.get("level", ""), es.get("level", ""))
    if es_value is not None:
        lines.append(f"- 指标：{es_type} = {es_value:.4f}（{es_level}效应）")
    lines.append("")

    # 6 结果解读与建议
    lines.append("**步骤六 · 结论与建议**")
    lines.append(f"> {interpretation}")
    lines.append("")

    if suggestions:
        for s in suggestions:
            lines.append(f"- {s}")
        lines.append("")

    return lines


def _spc_process_md(a: dict[str, Any], index: int, data_summary: dict[str, Any] | None) -> list[str]:
    """Return markdown lines for one SPC analysis's 5-step process."""
    lines: list[str] = []
    method_name = a.get("method_name", a.get("method", "SPC 控制图"))
    interpretation = a.get("interpretation", "")
    suggestions = a.get("suggestions", [])
    vis = (a.get("visualizations") or [{}])[0] if a.get("visualizations") else {}
    chart_data = vis.get("data") or {}
    points = chart_data.get("points", [])
    ucl = chart_data.get("ucl")
    cl = chart_data.get("cl")
    lcl = chart_data.get("lcl")
    chart_type = chart_data.get("chart_type", "")
    chart_title = vis.get("title", "")

    # Extract Y variable from title like "IX-MR 控制图 — Transactional CT"
    y_var = chart_title.split("—")[-1].strip() if "—" in chart_title else "—"

    anomaly_pts = [p for p in points if p.get("is_anomaly")]
    violated_rules: set[int] = set()
    for p in anomaly_pts:
        for r in (p.get("rule_violated") or []):
            violated_rules.add(r)

    rule_desc = {
        1: "1点超出3σ控制限",
        2: "连续9点落在中心线同侧",
        3: "连续6点递增或递减",
        4: "连续14点交替上下",
        5: "3点中有2点超出2σ",
        6: "5点中有4点超出1σ",
        7: "连续15点在1σ范围内",
        8: "连续8点在1σ范围外",
    }

    lines.append(f"### 分析 {index}：{method_name}")
    lines.append("")

    # 1 数据识别
    lines.append("**步骤一 · 数据识别**")
    lines.append(f"- Y 变量：{y_var}")
    lines.append(f"- 样本量：{len(points)} 个数据点")
    lines.append("- 数据类型：连续型（计量数据）")
    if data_summary:
        lines.append(f"- 原始数据集：{data_summary.get('rows', '-')} 行 × {data_summary.get('columns', '-')} 列")
    lines.append("")

    # 2 图型选择
    lines.append("**步骤二 · 图型选择**")
    lines.append(f"- 控制图类型：{chart_type}")
    type_reasons = {
        "IX-MR": "单值数据（子组大小=1），适用于个别值与移动极差控制图",
        "Xbar-R": "子组数据（子组大小2-9），适用于均值-极差控制图",
        "Xbar-S": "子组数据（子组大小≥10），适用于均值-标准差控制图",
        "P": "不合格品率数据，适用于 P 控制图",
        "NP": "不合格品数数据，适用于 NP 控制图",
        "C": "缺陷数数据（固定检验单位），适用于 C 控制图",
        "U": "单位缺陷数数据（可变检验单位），适用于 U 控制图",
    }
    lines.append(f"- 选型依据：{type_reasons.get(chart_type, f'根据数据特征选择 {chart_type} 控制图')}")
    lines.append("")

    # 3 控制限计算
    lines.append("**步骤三 · 控制限计算**")
    if ucl is not None:
        lines.append(f"- UCL（上控制限）= {ucl:.2f}")
    if cl is not None:
        lines.append(f"- CL（中心线）= {cl:.2f}")
    if lcl is not None:
        lines.append(f"- LCL（下控制限）= {lcl:.2f}")
    if ucl is not None and cl is not None:
        sigma = (ucl - cl) / 3
        lines.append(f"- σ（标准差）= {sigma:.2f}")
    lines.append("")

    # 4 异常检测
    lines.append("**步骤四 · 异常检测（西联规则）**")
    lines.append(f"- 异常点数：{len(anomaly_pts)} / {len(points)}")
    if violated_rules:
        for r in sorted(violated_rules):
            lines.append(f"- 规则 {r}：{rule_desc.get(r, f'自定义规则 {r}')}")
    else:
        lines.append("- 未触发任何西联规则，过程数据表现正常")
    lines.append("")

    # 5 结论与建议
    lines.append("**步骤五 · 结论与建议**")
    if len(anomaly_pts) == 0:
        status = "受控"
    elif len(anomaly_pts) <= 2:
        status = "预警"
    else:
        status = "失控"
    lines.append(f"- 过程状态：**{status}**")
    lines.append(f"> {interpretation}")
    lines.append("")
    if suggestions:
        for s in suggestions:
            lines.append(f"- {s}")
        lines.append("")

    return lines


def _safe_title(title: str | None, fallback: str) -> str:
    if title and str(title).strip():
        return str(title).strip()
    return fallback


def _render_chart_png(config: dict[str, Any]) -> bytes | None:
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm

        # Configure Chinese font support
        _zh_fonts = ["Microsoft YaHei", "SimHei", "WenQuanYi Micro Hei", "Noto Sans CJK SC", "Arial Unicode MS"]
        _found_font = None
        for fname in _zh_fonts:
            matches = fm.findfont(fname, fallback_to_default=False)
            if matches and "LastResort" not in matches:
                _found_font = fname
                break
        if _found_font:
            plt.rcParams["font.sans-serif"] = [_found_font] + plt.rcParams.get("font.sans-serif", [])
        plt.rcParams["axes.unicode_minus"] = False
    except Exception:
        return None

    chart_type = config.get("type")
    title = _safe_title(config.get("title"), "Chart")
    data = config.get("data") or {}

    fig, ax = plt.subplots(figsize=(6, 3.5))
    ax.set_title(title)

    def _maybe_sample(points: list[list[float]], max_n: int = 4000) -> list[list[float]]:
        if len(points) <= max_n:
            return points
        step = max(1, len(points) // max_n)
        return points[::step]

    try:
        if chart_type in {"scatter", "residual"}:
            points = data.get("points", [])
            points = _maybe_sample(points)
            xs = [p[0] for p in points]
            ys = [p[1] for p in points]
            ax.scatter(xs, ys, s=12, alpha=0.7)
            if config.get("xLabel"):
                ax.set_xlabel(str(config.get("xLabel")))
            if config.get("yLabel"):
                ax.set_ylabel(str(config.get("yLabel")))

        elif chart_type == "box":
            groups = data.get("groups", [])
            values = data.get("values", [])
            if groups and values:
                ax.boxplot(values, labels=[str(g) for g in groups], showfliers=False)
                if config.get("xLabel"):
                    ax.set_xlabel(str(config.get("xLabel")))
                if config.get("yLabel"):
                    ax.set_ylabel(str(config.get("yLabel")))

        elif chart_type == "bar":
            table = data.get("table")
            if table and isinstance(table, dict):
                series_names = list(table.keys())
                row_set = set()
                for col in series_names:
                    rows = table.get(col, {}) or {}
                    row_set.update(rows.keys())
                categories = list(row_set)
                bottoms = [0] * len(categories)
                for col in series_names:
                    vals = [int((table.get(col, {}) or {}).get(r, 0)) for r in categories]
                    ax.bar(categories, vals, bottom=bottoms, label=str(col))
                    bottoms = [b + v for b, v in zip(bottoms, vals)]
                ax.legend(fontsize=7)
            else:
                categories = data.get("categories", [])
                values = data.get("values", [])
                ax.bar(categories, values)
            if config.get("xLabel"):
                ax.set_xlabel(str(config.get("xLabel")))
            if config.get("yLabel"):
                ax.set_ylabel(str(config.get("yLabel")))

        elif chart_type == "distribution":
            bins = data.get("bins", [])
            counts = data.get("counts", [])
            if bins and counts:
                labels = [f"{b[0]}-{b[1]}" if isinstance(b, list) else str(b) for b in bins]
                ax.bar(labels, counts)
                ax.tick_params(axis="x", labelrotation=45)
            if config.get("xLabel"):
                ax.set_xlabel(str(config.get("xLabel")))
            if config.get("yLabel"):
                ax.set_ylabel(str(config.get("yLabel")))

        elif chart_type == "control_chart":
            points = data.get("points", [])
            ucl = data.get("ucl")
            cl = data.get("cl")
            lcl = data.get("lcl")
            chart_subtype = data.get("chart_type", "")

            xs = [p.get("x", i) for i, p in enumerate(points)]
            ys = [p.get("y", 0) for p in points]

            # Normal points line
            ax.plot(xs, ys, marker="o", markersize=4, linewidth=1, color="#3498db", label="数据", zorder=2)

            # Anomaly points
            anom_xs = [p.get("x", i) for i, p in enumerate(points) if p.get("is_anomaly")]
            anom_ys = [p.get("y", 0) for p in points if p.get("is_anomaly")]
            if anom_xs:
                ax.scatter(anom_xs, anom_ys, color="#e74c3c", s=50, zorder=3, label="异常点")

            # Control limits
            if ucl is not None:
                ax.axhline(y=ucl, color="#e74c3c", linestyle="--", linewidth=1.2, label=f"UCL={ucl:.2f}")
            if cl is not None:
                ax.axhline(y=cl, color="#2ecc71", linestyle="-", linewidth=1.5, label=f"CL={cl:.2f}")
            if lcl is not None:
                ax.axhline(y=lcl, color="#e74c3c", linestyle="--", linewidth=1.2, label=f"LCL={lcl:.2f}")

            ax.set_xlabel("样本序号")
            ax.set_ylabel(chart_subtype or "值")
            ax.legend(fontsize=7, loc="upper right")

        else:
            ax.text(0.5, 0.5, f"Unsupported chart type: {chart_type}", ha="center", va="center")
    except Exception:
        plt.close(fig)
        return None

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


# =========================================================================
# Markdown
# =========================================================================

def build_markdown_report(
    session: SessionModel,
    include_charts: bool,
    llm_config: dict | None = None,
) -> tuple[str, list[tuple[str, bytes]]]:
    """Return (markdown_text, [(image_filename, png_bytes), ...])."""
    lines: list[str] = []
    images: list[tuple[str, bytes]] = []

    lines.append("# 分析报告")
    lines.append("")
    if session.file_name:
        lines.append(f"> 数据文件：{session.file_name}　|　生成时间：{_to_iso(session.created_at)}")
        lines.append("")

    analyses = _collect_analyses(session)

    # ── 分析结论 ──
    if session.report_conclusion:
        lines.append(session.report_conclusion)
        lines.append("")
    else:
        lines.append("（未生成分析报告结论）")
        lines.append("")

    # ── 图表 ──
    if include_charts and analyses:
        lines.append("---")
        lines.append("")
        lines.append("## 图表")
        lines.append("")
        img_idx = 0
        for i, a in enumerate(analyses, 1):
            charts = a.get("visualizations") or []
            if not charts:
                continue
            if len(analyses) > 1:
                lines.append(f"### 分析 {i}")
                lines.append("")
            for j, chart in enumerate(charts, 1):
                title = _safe_title(chart.get("title"), f"图表 {i}.{j}")
                lines.append(f"**{title}**")
                png = _render_chart_png(chart)
                if png:
                    img_idx += 1
                    img_name = f"chart-{img_idx}.png"
                    images.append((img_name, png))
                    lines.append(f"![{title}]({img_name})")
                else:
                    lines.append("（图表渲染失败）")
                lines.append("")

    return "\n".join(lines).strip() + "\n", images


# =========================================================================
# DOCX
# =========================================================================

def _add_step(doc: Document, title: str, content: str, bold_content: str | None = None):
    """Add a step paragraph with bold title."""
    p = doc.add_paragraph()
    run_title = p.add_run(title + "\n")
    run_title.bold = True
    run_title.font.size = Pt(10)
    if bold_content:
        run_b = p.add_run(bold_content)
        run_b.bold = True
    else:
        p.add_run(content)


def build_docx_report(
    session: SessionModel,
    include_charts: bool,
    llm_config: dict | None = None,
) -> bytes:
    doc = Document()

    doc.add_heading("分析报告", level=0)
    if session.file_name:
        doc.add_paragraph(f"数据文件：{session.file_name}　|　生成时间：{_to_iso(session.created_at)}")

    analyses = _collect_analyses(session)

    # ── 分析结论 ──
    if session.report_conclusion:
        for para in session.report_conclusion.split("\n"):
            stripped = para.strip()
            if stripped:
                doc.add_paragraph(stripped)
    else:
        doc.add_paragraph("（未生成分析报告结论）")

    # ── 图表 ──
    if include_charts and analyses:
        doc.add_heading("图表", level=1)
        for i, a in enumerate(analyses, 1):
            charts = a.get("visualizations") or []
            if not charts:
                continue
            if len(analyses) > 1:
                doc.add_heading(f"分析 {i}", level=2)
            for j, chart in enumerate(charts, 1):
                title = _safe_title(chart.get("title"), f"图表 {i}.{j}")
                doc.add_paragraph(title)
                png = _render_chart_png(chart)
                if png:
                    doc.add_picture(io.BytesIO(png), width=Pt(420))
                else:
                    doc.add_paragraph("图表渲染失败。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
